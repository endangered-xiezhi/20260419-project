from pathlib import Path
import shutil
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "artifacts" / "demo"
FONT = "Source Han Sans SC"
HEADING_FONT = "Source Han Sans SC"
NAVY = "0B2545"
BLUE = "176B87"
PALE = "EAF3F6"
GRAY = "F2F4F7"
MUTED = "667085"
RED = "B42318"
GOLD = "8A6116"
GREEN = "067647"
DEMO_COMPANY = "XXX股份有限公司"


def set_run(run, size=11, bold=False, color="172B4D", font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_mar = cell._tc.get_or_add_tcPr().find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                cell._tc.get_or_add_tcPr().append(tc_mar)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                item = OxmlElement(f"w:{side}")
                item.set(qn("w:w"), str(value))
                item.set(qn("w:type"), "dxa")
                tc_mar.append(item)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run(run, size=9, color=MUTED)


def setup_doc(title, subtitle, doc_type):
    doc = Document()
    section = doc.sections[0]
    # Named override to standard_business_brief: A4 + Chinese corporate font.
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, NAVY),
    ):
        style = doc.styles[style_name]
        style.font.name = HEADING_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run(f"三会治理数字化演示  |  {doc_type}")
    set_run(r, size=9, bold=True, color=MUTED, font=HEADING_FONT)
    add_page_field(section.footer.paragraphs[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run(r, size=23, bold=True, color=NAVY, font=HEADING_FONT)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(subtitle)
    set_run(r, size=12.5, color=MUTED, font=HEADING_FONT)
    return doc


def add_metadata(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        shade(cells[0], PALE)
        p = cells[0].paragraphs[0]
        r = p.add_run(label)
        set_run(r, size=10, bold=True, color=NAVY, font=HEADING_FONT)
        p = cells[1].paragraphs[0]
        r = p.add_run(value)
        set_run(r, size=10)
    set_table_widths(table, [1900, 7460])
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size={1: 16, 2: 13, 3: 12}[level], bold=True,
            color=BLUE if level < 3 else NAVY, font=HEADING_FONT)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.3
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run(r, bold=True, color=NAVY, font=HEADING_FONT)
        r = p.add_run(text[len(bold_lead):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)


def add_bullet(doc, text, color="172B4D"):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run(r, color=color)


def add_callout(doc, title, text, fill=PALE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [9360])
    shade(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run(f"{title}\n")
    set_run(r, size=11, bold=True, color=color, font=HEADING_FONT)
    r = p.add_run(text)
    set_run(r, size=10.5, color=color)
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        shade(table.rows[0].cells[i], GRAY)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run(r, size=9.5, bold=True, color=NAVY, font=HEADING_FONT)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            set_run(r, size=9.3)
    set_table_widths(table, widths)
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def build_meeting_demo():
    doc = setup_doc(
        "2026年度第一次临时股东会全流程演示档案",
        "从飞书会议数据、妙记纪要、议案生成、表决统计到合规归档的完整样例",
        "股东会全流程档案",
    )
    add_metadata(doc, [
        ("演示公司", f"{DEMO_COMPANY}（演示）"),
        ("会议编号", "GMS-2026-001"),
        ("会议时间", "2026年7月10日 14:00—16:20"),
        ("会议地点", "上海市浦东新区治理路88号第一会议室"),
        ("召开方式", "现场会议与飞书视频会议相结合"),
        ("数据来源", "飞书公司主体表、会议表、股东表、议案表、表决表及飞书妙记"),
        ("档案状态", "演示定稿｜已完成合规复核｜待电子签章"),
    ])
    add_callout(
        doc,
        "演示结论",
        "系统已自动识别会议类型为“临时股东会”，读取公司主体名称、会议日期、参会股东、持股数量、妙记摘要和三项关联议案；对缺失的授权委托原件作出提示，并形成完整会议审查记录。",
    )
    add_heading(doc, "一、会议概览", 1)
    add_body(doc, "本次会议由董事会依法召集，董事长王明远主持。会议审议年度治理、利润分配和三会治理数字化三项议案。系统以会议表记录为主键，联动公司主体表、股东表、人员表、议案表和表决表，避免重复录入。")
    add_matrix(doc, ["项目", "应到/总数", "实到/代表", "比例", "系统判断"], [
        ["股东人数", "5名", "5名", "100%", "达到召开条件"],
        ["有表决权股份", "10,000,000股", "10,000,000股", "100%", "达到表决基础"],
        ["关联议案", "3项", "3项", "100%", "资料齐备"],
        ["会前通知", "提前18日", "已送达5名股东", "100%", "期限满足演示规则"],
    ], [1500, 1800, 2000, 1400, 2660])

    add_heading(doc, "二、参会股东及表决权", 1)
    add_matrix(doc, ["股东", "股东性质", "持股数量", "持股比例", "出席方式", "授权状态"], [
        ["上海智源产业投资有限公司", "法人股东", "4,000,000股", "40.00%", "现场", "法定代表人出席"],
        ["李四", "自然人股东", "1,820,000股", "18.20%", "现场", "本人出席"],
        ["启航创新基金合伙企业", "合伙企业", "1,500,000股", "15.00%", "视频", "授权代表出席"],
        ["周岚", "自然人股东", "1,380,000股", "13.80%", "现场", "本人出席"],
        ["海岳员工持股平台", "合伙企业", "1,300,000股", "13.00%", "视频", "执行事务合伙人委派"],
    ], [1900, 1300, 1500, 1200, 1200, 2260])

    add_heading(doc, "三、会议议程", 1)
    for text in [
        "14:00—14:10　主持人宣布会议开始，核验股东及授权代表身份，报告出席情况。",
        "14:10—14:25　董事会秘书说明会议通知、召集程序、计票与监票安排。",
        "14:25—15:10　逐项宣读三项议案，相关负责人报告背景、影响和执行方案。",
        "15:10—15:40　股东提问，董事长、财务负责人和董事会秘书现场答复。",
        "15:40—16:00　股东填写表决票，监票人、计票人核验并统计真实表决意见。",
        "16:00—16:15　宣读表决结果和会议决议。",
        "16:15—16:20　主持人宣布会议结束，系统启动合规审查和归档流程。",
    ]:
        add_bullet(doc, text)

    page_break(doc)
    add_heading(doc, "四、议案一：2025年度董事会工作报告", 1)
    add_body(doc, "议案背景：2025年度，公司围绕主营业务提质、研发成果转化和治理体系规范化开展经营。董事会共召开8次会议，审议重大经营事项31项，所有会议均形成完整通知、议案、表决和决议档案。")
    add_body(doc, "议案内容：董事会提请股东会确认2025年度工作报告。报告覆盖战略执行、经营结果、风险管理、内部控制、重大合同、关联交易识别、投资者权益保护和董事履职评价。")
    add_body(doc, "执行建议：股东会审议通过后，由董事会秘书将工作报告、股东提问答复和表决结果一并归档；报告提出的六项治理改进任务纳入2026年度督办清单。")
    add_heading(doc, "五、议案二：2025年度利润分配方案", 1)
    add_body(doc, "财务基础：经演示审计口径测算，公司2025年度实现归属于母公司股东的净利润12,680,000元，期末可供分配利润28,450,000元，经营活动现金流量净额18,920,000元。")
    add_body(doc, "分配方案：拟以总股本10,000,000股为基数，每10股派发现金红利3.00元（含税），合计派发现金红利3,000,000元；不送红股，不以资本公积转增股本。")
    add_body(doc, "风险控制：实施前由财务负责人复核可分配利润、资金安排和税务处理；如股本发生变化，保持分配总额不变并依法调整每股分配比例。")
    add_heading(doc, "六、议案三：三会治理数字化建设方案", 1)
    add_body(doc, "建设目的：统一公司主体、会议、人员、议案、文书和表决数据口径，完善通知送达、电子会议、关联事项回避、会议资料电子归档和权限管理。方案遵循“权责明确、程序可验证、档案可追溯”的原则。")
    add_matrix(doc, ["建设主题", "现状摘要", "拟建设内容", "审查关注"], [
        ["会议通知", "以人工送达为主", "增加飞书、电子邮件等可验证方式", "保留送达记录"],
        ["电子会议", "记录分散", "视频会议与身份核验记录关联", "确保同步参与"],
        ["关联回避", "人工统计", "记录申报、回避、票数扣除和复核", "防止计票错误"],
        ["电子档案", "纸质与电子分散", "电子原件与签章文件同步归档", "权限与保存期限"],
    ], [1500, 2200, 3300, 2360])

    add_heading(doc, "七、飞书妙记摘录与关键问答", 1)
    add_callout(doc, "妙记摘要（15:18）", "股东李四询问数字化系统是否会改变董事会授权。董事长答复：系统只记录和校验既有权限，不扩大任何机构的法定或内部授权。董事会秘书进一步说明，系统会对授权事项设置金额、期限和事项类型限制。")
    add_body(doc, "财务负责人就利润分配后的现金流安全边界作出说明：分配完成后，公司仍保留覆盖未来十二个月经营预算及已批准资本性支出的资金，且不影响到期债务偿付。")
    add_body(doc, "监票人确认：系统生成的空白表决票只在“文书表”登记，不自动形成表决意见；只有签署后的真实选择被复核录入后，才在“表决表”创建同意、反对或弃权记录。")

    page_break(doc)
    add_heading(doc, "八、表决结果", 1)
    add_matrix(doc, ["议案", "同意", "反对", "弃权", "回避", "结果"], [
        ["董事会工作报告", "10,000,000股（100%）", "0股", "0股", "0股", "通过"],
        ["利润分配方案", "8,500,000股（85%）", "1,500,000股（15%）", "0股", "0股", "通过"],
        ["三会治理数字化建设", "10,000,000股（100%）", "0股", "0股", "0股", "通过"],
    ], [2500, 1900, 1400, 1200, 1100, 1260])
    add_callout(doc, "数据分层规则", "空白表决票：写入“文书表”，生成状态为“已生成”；真实表决意见：经签署或电子确认后写入“表决表”。两类记录永久分开，防止将文书生成误认为已经投票。", fill="FFF7E6", color=GOLD)

    add_heading(doc, "九、会议决议", 1)
    add_body(doc, "经出席会议并具有表决权的股东表决，本次会议审议的三项议案均获通过。")
    add_body(doc, "会议授权董事会及其授权人员办理利润分配实施、数字化建设和会议档案归档事项。授权不得改变股东会决议的实质内容；执行中如发生重大变化，应重新履行相应决策程序。")

    add_heading(doc, "十、合规审查摘要", 1)
    add_matrix(doc, ["审查维度", "核验结果", "证据来源", "结论"], [
        ["召集权限", "董事会召集、董事长主持", "会议表、董事会决议", "通过"],
        ["通知期限", "提前18日完成通知", "通知表、送达回执", "通过"],
        ["出席与身份", "5名股东全部核验", "股东表、授权文件", "通过"],
        ["议案完整性", "背景、正文、依据、建议齐备", "会议表、议案表", "通过"],
        ["表决统计", "空白票与真实意见分离", "文书表、表决表", "通过"],
        ["档案完整性", "授权代表原件待补扫", "归档清单", "中风险提示"],
    ], [1700, 2700, 2400, 2560])
    add_callout(doc, "AI 审查结论｜合规指数 92/100", "会议召集、通知、出席、议案审议和表决统计流程总体完整。唯一待办事项为两份授权代表文件的纸质原件扫描件尚未进入最终档案包；建议在决议签署前完成补扫并由董事会秘书复核。", fill="ECFDF3", color=GREEN)

    add_heading(doc, "十一、归档清单", 1)
    for text in [
        "会议通知及5份可验证送达记录；",
        "三项议案、说明材料及数字化建设方案；",
        "股东名册、持股数量快照、授权委托文件及身份核验记录；",
        "飞书妙记链接、逐字稿、会议纪要和关键问答摘要；",
        "15份空白表决票生成记录、15份真实表决意见记录及表决统计表；",
        "会议决议、签字页和合规审查报告；",
        "文书包哈希值、归档时间、归档责任人及访问权限记录。",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "系统归档标识", "会议记录 ID：rec_demo_gms_2026001｜审批实例：APPROVAL-DEMO-20260710-001｜归档包：GMS-2026-001-v1.0.zip｜生成时间：2026年7月19日 16:30（演示数据）")

    path = OUT / "三会治理全流程演示样例.docx"
    doc.save(path)
    return path


def build_review_demo():
    doc = setup_doc(
        "股东会合规 AI 审查报告",
        "基于飞书公司主体表、会议表、议案表、妙记和表决记录的多源交叉审查样例",
        "合规审查报告",
    )
    add_metadata(doc, [
        ("审查对象", f"{DEMO_COMPANY}2026年度第一次临时股东会（演示）"),
        ("公司主体", f"{DEMO_COMPANY}｜股份有限公司"),
        ("会议记录", "飞书会议表：GMS-2026-001"),
        ("妙记来源", "2026年7月10日临时股东会飞书妙记"),
        ("审查引擎", "会议合规 AI 审查工作流 v2.3（演示）"),
        ("审查时间", "2026年7月19日 16:35"),
        ("总体评分", "92 / 100｜总体可通过，1项中风险待闭环"),
    ])
    add_callout(doc, "审查方法", "AI 不仅检查单一文书，而是以会议记录为主线，将公司主体、股东名册、通知送达、议案正文、妙记发言、空白表决票和真实表决意见进行交叉核验。缺失信息单独列入待办，不用占位符伪装成已完成。")

    add_heading(doc, "一、审查范围与数据可信度", 1)
    add_matrix(doc, ["数据源", "读取内容", "完整度", "可信度判断"], [
        ["飞书会议表", "会议类型、日期、地点、召集人、参会人员、状态", "100%", "结构化主数据"],
        ["公司主体表", "公司全称、主体类型和登记信息", "100%", "结构化主数据"],
        ["飞书妙记", "关键发言、提问、答复、时间点", "96%", "需与正式纪要确认"],
        ["议案表", "三项议案正文、依据、建议和关联会议", "100%", "结构化审议材料"],
        ["文书表/表决表", "空白票与真实意见分层记录", "100%", "可追溯业务记录"],
        ["授权文件", "2份电子件、2份纸质原件待补扫", "75%", "存在归档缺口"],
    ], [1700, 3500, 1100, 3060])

    add_heading(doc, "二、会议规则核验", 1)
    add_body(doc, "系统根据会议类型和已配置的会议治理规则，对本次临时股东会的召集、通知、表决和归档信息进行结构化核验。")
    add_matrix(doc, ["规则主题", "核验要求", "会议实际", "匹配结果"], [
        ["召集主体", "召集主体信息完整并有记录", "董事会决议召集", "一致"],
        ["通知内容", "载明时间、地点、方式和审议事项", "通知含全部必要项目", "一致"],
        ["通知期限", "满足系统配置的通知期限", "提前18日完成送达", "一致"],
        ["表决基础", "票权记录清晰，回避单独处理", "按股份数统计，无回避事项", "一致"],
        ["会议记录", "记录出席、审议、发言要点和表决结果", "妙记+正式纪要齐备", "一致"],
        ["档案管理", "授权文件、表决票和决议一并归档", "纸质原件扫描待补", "部分一致"],
    ], [1600, 3600, 2300, 1860])

    page_break(doc)
    add_heading(doc, "三、逐项审查结论", 1)
    add_heading(doc, "3.1 会议召集与通知", 2)
    add_body(doc, "结论：通过。会议由董事会召集，会议通知在会议召开18日前向全部5名股东发送。通知载明日期、时间、地点、召开方式、审议议案、联系人和表决方式；送达记录可验证。")
    add_heading(doc, "3.2 出席、授权与表决权", 2)
    add_body(doc, "结论：有条件通过。5名股东均由本人、法定代表人或授权代表出席，代表全部有表决权股份。两名授权代表的电子授权文件已核验，但纸质原件扫描尚未归档，不影响演示计票，但构成档案完整性待办。")
    add_heading(doc, "3.3 议案完整性与权限边界", 2)
    add_body(doc, "结论：通过。三项议案均与会议表正确关联，包含议案背景、具体方案、执行建议和适用规则。利润分配由股东会作出最终决定，未发现系统擅自改变会议权限的情形。")
    add_heading(doc, "3.4 妙记与正式文书一致性", 2)
    add_body(doc, "结论：通过。妙记中关于授权边界、现金流安全和表决票性质的关键答复，均已进入正式会议纪要。系统未将未确认的口语表达直接写入决议，仅作为问答记录保留。")
    add_heading(doc, "3.5 表决真实性与统计", 2)
    add_body(doc, "结论：通过。文书表中的15份记录均标记为“空白表决票”，不含表决意见；表决表中的15份记录均有股东、议案、票权、真实意见和表决时间。系统按议案汇总票数并与出席股份总数进行勾稽，差额为零。")

    add_heading(doc, "四、风险清单与整改闭环", 1)
    add_matrix(doc, ["编号", "风险", "等级", "证据", "整改措施", "责任人/期限"], [
        ["R-01", "授权代表纸质原件扫描未归档", "中", "归档清单缺2份扫描件", "补扫、核验签章并关联会议记录", "董事会秘书/2个工作日"],
        ["R-02", "妙记为辅助证据，需固定最终版本", "低", "妙记可持续编辑", "导出逐字稿并记录版本哈希", "会议记录人/当日"],
        ["R-03", "数字化建设任务尚未分解", "低", "决议已授权但未形成任务清单", "建立实施任务并回填结果", "项目负责人/10日内"],
    ], [800, 2150, 800, 1900, 2450, 1260])
    add_callout(doc, "阻断判断", "R-01 不阻断本次演示会议决议生成，但在最终归档状态从“待归档”变更为“已归档”之前必须完成。若授权文件内容与已核验电子件不一致，应立即暂停归档并重新核验表决权。", fill="FFF4ED", color=RED)

    add_heading(doc, "五、AI 交叉核验记录", 1)
    add_matrix(doc, ["核验关系", "系统比对", "结果"], [
        ["会议表 ↔ 议案表", "会议 ID 与三项议案关联字段一致", "通过"],
        ["股东表 ↔ 表决表", "每名股东每项议案仅一条有效真实意见", "通过"],
        ["文书表 ↔ 表决表", "空白票生成记录未写入表决意见", "通过"],
        ["妙记 ↔ 会议纪要", "三段关键问答均进入正式纪要", "通过"],
        ["公司主体表 ↔ 文书", f"全部文书使用“{DEMO_COMPANY}”", "通过"],
        ["通知表 ↔ 会议日期", "送达日至召开日间隔18日", "通过"],
        ["出席股份 ↔ 计票合计", "10,000,000股 = 有效票合计", "通过"],
    ], [2300, 5100, 1960])

    add_heading(doc, "六、审查结果摘要", 1)
    add_metadata(doc, [
        ("审查事项", "2026年度第一次临时股东会合规审查"),
        ("负责部门", "董事会办公室"),
        ("合规评分", "92分"),
        ("审查结论", "总体可通过，授权文件纸质原件扫描待补"),
        ("处理建议", "法务复核 → 董事会秘书补齐档案 → 归档管理员确认入库"),
        ("关联资料", "会议记录、妙记、议案包、表决统计和审查报告"),
    ])
    add_body(doc, "审查完成后，将结论和整改事项写入会议审查记录；整改完成后更新“审查状态”为“已通过”，未完成则保持“需整改”并保留处理意见。")

    add_heading(doc, "七、最终意见", 1)
    add_callout(doc, "总体意见｜建议有条件通过", "本次会议的召集、通知、出席、议案、表决和决议链条完整，系统记录之间可以相互印证。建议完成R-01整改后将档案状态更新为“已归档”，并按期关闭数字化建设任务。", fill="ECFDF3", color=GREEN)
    add_body(doc, "本报告为产品演示样例，用于展示公司主体识别、多源数据读取、证据交叉核验和风险分级能力。正式项目中，应使用公司实际会议制度及适用法律法规，并由有权人员完成最终复核。")

    path = OUT / "会议合规AI审查报告样例.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    web_demo = Path(__file__).resolve().parents[1] / "public" / "demo"
    web_demo.mkdir(parents=True, exist_ok=True)
    for generated in (build_meeting_demo(), build_review_demo()):
        shutil.copy2(generated, web_demo / generated.name)
        print(generated)
